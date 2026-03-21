from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class Exporter:
    """导出器，支持导出为PDF和Word格式"""
    
    def __init__(self):
        self.export_folder = 'exports'
        os.makedirs(self.export_folder, exist_ok=True)
        
        # 注册中文字体（如果需要）
        # pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
    
    def export_to_pdf(self, record):
        """导出为PDF"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"analysis_{record['id']}_{timestamp}.pdf"
        filepath = os.path.join(self.export_folder, filename)
        
        try:
            c = canvas.Canvas(filepath, pagesize=A4)
            width, height = A4
            
            # 标题
            c.setFont("Helvetica-Bold", 16)
            c.drawString(inch, height - inch, f"Analysis Report: {record['title']}")
            
            y = height - 1.5 * inch
            
            # 统计信息
            c.setFont("Helvetica-Bold", 12)
            c.drawString(inch, y, "Statistics:")
            y -= 0.3 * inch
            
            c.setFont("Helvetica", 10)
            stats = record['statistics']
            c.drawString(inch, y, f"Word Count: {stats.get('word_count', 0)}")
            y -= 0.2 * inch
            c.drawString(inch, y, f"Key Points: {stats.get('keypoint_count', 0)}")
            y -= 0.3 * inch
            
            # 关键点
            c.setFont("Helvetica-Bold", 12)
            c.drawString(inch, y, "Key Points:")
            y -= 0.3 * inch
            
            c.setFont("Helvetica", 10)
            for idx, kp in enumerate(record['keypoints'][:10], 1):
                if y < inch:
                    c.showPage()
                    y = height - inch
                
                text = f"{idx}. {kp['content'][:100]}..."
                c.drawString(inch, y, text)
                y -= 0.3 * inch
            
            c.save()
            return filepath
        
        except Exception as e:
            print(f"PDF导出错误: {str(e)}")
            # 简化版本导出
            return self._export_simple_pdf(record)
    
    def _export_simple_pdf(self, record):
        """简化的PDF导出（不使用中文字体）"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"analysis_{record['id']}_{timestamp}.pdf"
        filepath = os.path.join(self.export_folder, filename)
        
        c = canvas.Canvas(filepath, pagesize=A4)
        width, height = A4
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, "Analysis Report")
        
        y = height - 1.5 * inch
        c.setFont("Helvetica", 12)
        c.drawString(inch, y, f"Record ID: {record['id']}")
        y -= 0.3 * inch
        c.drawString(inch, y, f"Filename: {record['filename']}")
        y -= 0.3 * inch
        
        stats = record['statistics']
        c.drawString(inch, y, f"Word Count: {stats.get('word_count', 0)}")
        y -= 0.3 * inch
        c.drawString(inch, y, f"Key Points: {stats.get('keypoint_count', 0)}")
        
        c.save()
        return filepath
    
    def export_to_docx(self, record):
        """导出为Word文档"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"analysis_{record['id']}_{timestamp}.docx"
        filepath = os.path.join(self.export_folder, filename)
        
        try:
            doc = Document()
            
            # 标题
            title = doc.add_heading(record['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 文档信息
            doc.add_heading('文档信息', 1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_table.rows[0].cells[0].text = '文件名'
            info_table.rows[0].cells[1].text = record['filename']
            info_table.rows[1].cells[0].text = '分析时间'
            info_table.rows[1].cells[1].text = record.get('created_at', '')
            info_table.rows[2].cells[0].text = '字数'
            info_table.rows[2].cells[1].text = str(record['statistics'].get('word_count', 0))
            
            # 统计信息
            doc.add_heading('统计信息', 1)
            stats = record['statistics']
            
            p = doc.add_paragraph()
            p.add_run(f"关键点数量：{stats.get('keypoint_count', 0)}\n").bold = True
            p.add_run(f"平均重要性：{stats.get('avg_importance', 0)}\n")
            
            if 'top_keywords' in stats:
                p.add_run(f"关键词：{', '.join(stats['top_keywords'][:10])}\n")
            
            # 关键点列表
            doc.add_heading('关键点详情', 1)
            
            for idx, kp in enumerate(record['keypoints'], 1):
                # 关键点标题
                kp_heading = doc.add_heading(f"关键点 {idx}", 2)
                
                # 关键点内容
                p = doc.add_paragraph()
                
                # 类别和重要性
                run = p.add_run(f"[{kp['category']}] ")
                run.font.color.rgb = RGBColor(0, 112, 192)
                run.bold = True
                
                run = p.add_run(f"重要性：{kp['importance']}/100\n")
                run.font.color.rgb = RGBColor(255, 0, 0) if kp['importance'] > 80 else RGBColor(0, 0, 0)
                
                # 内容
                p.add_run(f"{kp['content']}\n")
                
                # 注释
                if 'annotation' in kp:
                    run = p.add_run(f"💡 {kp['annotation']}\n")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
                
                doc.add_paragraph()  # 空行
            
            # 摘要报告
            doc.add_heading('摘要报告', 1)
            summary = record['summary']
            
            if 'core_points' in summary and summary['core_points']:
                doc.add_heading('核心观点', 2)
                for point in summary['core_points']:
                    doc.add_paragraph(point, style='List Bullet')
            
            if 'key_data' in summary and summary['key_data']:
                doc.add_heading('关键数据', 2)
                for data in summary['key_data']:
                    doc.add_paragraph(data, style='List Bullet')
            
            if 'conclusions' in summary and summary['conclusions']:
                doc.add_heading('结论总结', 2)
                for conclusion in summary['conclusions']:
                    doc.add_paragraph(conclusion, style='List Bullet')
            
            # 保存文档
            doc.save(filepath)
            return filepath
        
        except Exception as e:
            print(f"Word导出错误: {str(e)}")
            raise e
