import re
import unicodedata
import jieba
import jieba.analyse
from difflib import SequenceMatcher
from llm_service import LLMService
import fitz  # PyMuPDF — 统一文本搜索 + 原生高亮标注引擎
import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

class TextAnalyzer:
    """文本分析器，提取关键点、生成摘要和统计信息"""
    
    def __init__(self):
        # 初始化jieba
        jieba.setLogLevel(jieba.logging.INFO)
        # 初始化大模型LLM服务
        self.llm_service = LLMService()
        # 存储结构化数据（用于精准定位）
        self.current_structured_data = None
        print("LLM 服务已启用")
    
    def analyze(self, text, structured_data=None, api_provider='deepseek'):
        """
        分析文本，返回完整的分析结果
        两阶段处理：
          阶段1：提取图片信息 → 独立交给大模型分析
          阶段2：图片分析结果 + 主文本 → 一起送入大模型做最终分析
            
        Args:
            text: 待分析文本
            structured_data: 结构化数据（包含 PDF 的图片、表格、文件路径等信息）
            api_provider: API 提供商 ('deepseek' | 'qwen')
        """
        # 保存结构化数据
        self.current_structured_data = structured_data

        print(f"使用 {api_provider} 大模型API进行文本分析...")

        # 获取文件路径和文件大小
        file_path = structured_data.get('filepath') if structured_data else None
        file_size = structured_data.get('file_size', 0) if structured_data else 0

        # ── 阶段1：图片独立处理 ──────────────────────────────
        image_descriptions = ''
        if structured_data and structured_data.get('structured_content'):
            sc = structured_data['structured_content']
            image_items = [x for x in sc if x['type'] == 'image']
            table_count = len([x for x in sc if x['type'] == 'table'])

            if image_items:
                print(f"🖼️ 阶段1：独立分析 {len(image_items)} 张图片...")
                image_descriptions = self.llm_service.analyze_images(image_items, provider=api_provider)
                if image_descriptions:
                    print(f"✅ 图片独立分析完成，结果长度: {len(image_descriptions)} 字")

            if table_count > 0 or image_items:
                extra_info = f"\n\n[文档结构信息] 此文档包含 {table_count} 个表格和 {len(image_items)} 张图片。"
                text += extra_info

        # ── 阶段2：主文本 + 图片分析结果 → 最终分析 ─────────
        print(f"📝 阶段2：综合分析（文件直传 → 降级文本提取）...")
        llm_result = self.llm_service.analyze_text(
            text,
            provider=api_provider,
            file_path=file_path,
            file_size=file_size,
            image_descriptions=image_descriptions
        )

        if llm_result:
            return self._process_llm_result(text, llm_result)

        print("大模型分析失败，回退到传统方法")
        return self._traditional_analyze(text)
    
    def _traditional_analyze(self, text):
        """简单传统分析（降级方案）"""
        sentences = self._split_sentences(text)
        title = self._extract_title(sentences)
            
        # 简单选取前 5 句作为关键点
        keypoints = []
        for i, s in enumerate(sentences[:5]):
            keypoints.append({
                'id': i + 1,
                'content': s,
                'summary': s[:20],
                'importance': 70,
                'category': '重点内容',
                'position': i,
                'annotation': '自动提取的关键句',
                'color': '#ffd43b'
            })
                
        summary = {'core_points': [s for s in sentences[:3]], 'key_data': [], 'conclusions': []}
        statistics = self._generate_statistics(text, keypoints)
        highlights = self._generate_highlights(text, keypoints)
            
        return {
            'title': title,
            'keypoints': keypoints,
            'summary': summary,
            'statistics': statistics,
            'highlights': highlights
        }
    
    def _process_llm_result(self, text, llm_result):
        title = llm_result.get('title', '分析报告')
        
        raw_llm_keypoints = []
        for idx, arg in enumerate(llm_result.get('core_arguments', [])):
            point_text  = arg.get('point', '')
            evidence_list = arg.get('evidence', [])
            importance  = arg.get('importance', 95)
            point_page  = arg.get('point_page')
            point_ctx   = arg.get('point_context', '')
            ann_label   = arg.get('annotation_label', '\u6838心论点')

            # evidence 字符串展示
            evidence_texts = []
            evidence_data = []  # 结构化 evidence 数据（传给前端）
            for ev in evidence_list:
                if isinstance(ev, dict):
                    ev_text = ev.get('text', '')
                    evidence_texts.append(ev_text)
                    evidence_data.append({
                        'text': ev_text,
                        'page': ev.get('page'),
                        'context': ev.get('context', '')
                    })
                else:
                    evidence_texts.append(str(ev))
                    evidence_data.append({'text': str(ev), 'page': None, 'context': ''})

            evidence_str = "\n".join([f"• {ev}" for ev in evidence_texts])
            point_annotation = f"【{ann_label}】\n{point_text}\n\n【支撑论据】\n{evidence_str}"

            point_id = (idx + 1) * 100
            raw_llm_keypoints.append({
                'id':          point_id,
                'content':     point_text,
                'context':     point_ctx,
                'category':    '核心论点',
                'importance':  importance,
                'annotation':  point_annotation,
                'fixed_color': '#ff6b6b',
                'page':        point_page,
                'evidence':    evidence_data,  # 结构化论据数据（前端展开用）
                'annotation_label': ann_label,
                'rationale':   arg.get('rationale', ''),
            })

            # 每个论据独立高亮
            for e_idx, ev in enumerate(evidence_list):
                if isinstance(ev, dict):
                    ev_text = ev.get('text', '')
                    ev_page = ev.get('page')
                    ev_ctx  = ev.get('context', '')
                else:
                    ev_text = str(ev)
                    ev_page = None
                    ev_ctx  = ''

                raw_llm_keypoints.append({
                    'id':          point_id + e_idx + 1,
                    'content':     ev_text,
                    'context':     ev_ctx,
                    'category':    '支撑论据',
                    'importance':  max(30, importance - 40),
                    'annotation':  f"【支撑论据】针对论点：{point_text[:50]}...",
                    'fixed_color': '#51cf66',
                    'page':        ev_page,
                })
        
        # 匹配到原文
        matched_keypoints = self._match_keypoints_to_text(text, raw_llm_keypoints)
        
        summary = {
            'core_points': [arg.get('point', '') for arg in llm_result.get('core_arguments', [])],
            'key_data': llm_result.get('key_data', []),  # 结构化关键数据（label/value/context/page）
            'conclusions': [llm_result.get('summary', '')]
        }
        
        statistics = self._generate_statistics(text, matched_keypoints)
        highlights = self._generate_highlights(text, matched_keypoints)
        
        return {
            'title': title,
            'keypoints': matched_keypoints,
            'summary': summary,
            'statistics': statistics,
            'highlights': highlights
        }
    
    def _match_keypoints_to_text(self, text, llm_keypoints):
        """
        将大模型提取的关键点匹配到原文中
        对于PDF，使用页码信息进行精准定位
        """
        matched_keypoints = []
        sentences = self._split_sentences(text)
        
        # 获取文件路径（如果是 PDF 则用于精准定位）
        pdf_path = None
        if self.current_structured_data and self.current_structured_data.get('filepath'):
            filepath = self.current_structured_data['filepath']
            if filepath.lower().endswith('.pdf'):
                pdf_path = filepath
        
        for idx, llm_kp in enumerate(llm_keypoints):
            kp_content = llm_kp.get('content', '')
            kp_page = llm_kp.get('page')  # 获取页码提示
            
            # 方法1：PDF精准定位 - 如果有PDF路径和页码信息
            matched_text = None
            match_method = 'unknown'
            pdf_location = None
            
            if pdf_path and kp_page:
                pdf_location = self._locate_text_in_pdf(pdf_path, kp_content, kp_page)
                if pdf_location:
                    matched_text = kp_content
                    match_method = 'pdf_location'
            
            # 方法2：精确匹配 - 直接在原文中查找
            if not matched_text and kp_content in text:
                matched_text = kp_content
                match_method = 'exact'
            
            # 方法3：模糊匹配 - 找到最相似的句子
            if not matched_text:
                matched_text, match_method = self._fuzzy_match_sentence(kp_content, sentences)
            
            if matched_text:
                # 使用传入的 annotation
                final_annotation = llm_kp.get('annotation', '')
                
                # 生成摘要描述
                point_summary = matched_text[:30] + '...'
                
                # 优先使用 fixed_color，否则根据重要性动态生成
                final_color = llm_kp.get('fixed_color') or self._get_color_by_importance(llm_kp.get('importance', 70) / 100)
                
                keypoint = {
                    'id': llm_kp.get('id') or (idx + 1),
                    'content': matched_text,
                    'llm_content': kp_content,
                    'summary': point_summary,
                    'importance': llm_kp.get('importance', 70),
                    'category': llm_kp.get('category', '核心内容'),
                    'position': sentences.index(matched_text) if matched_text in sentences else -1,
                    'annotation': final_annotation,
                    'color': final_color,
                    'match_method': match_method,
                    'page': kp_page,
                    'pdf_location': pdf_location,
                    'evidence': llm_kp.get('evidence', []),           # 结构化论据（前端展开用）
                    'annotation_label': llm_kp.get('annotation_label', ''),  # 标注标签
                    'rationale': llm_kp.get('rationale', ''),         # 评分理由
                }
                matched_keypoints.append(keypoint)
        
        # 按重要性排序
        matched_keypoints.sort(key=lambda x: x['importance'], reverse=True)
        
        return matched_keypoints
    
    def _fuzzy_match_sentence(self, target_sentence, sentences):
        """
        模糊匹配句子：从原文句子中找到与目标句子最相似的
        使用多种策略：
        1. 序列相似度
        2. 关键词重叠度
        3. 长度相似度
        """
        if not sentences:
            return target_sentence, 'none'
        
        best_match = None
        best_score = 0.0
        match_method = 'fuzzy'
        
        # 提取目标句子的关键词
        target_keywords = set(jieba.cut(target_sentence))
        
        for sentence in sentences:
            # 计算多个相似度指标
            
            # 1. 序列相似度（SequenceMatcher）
            seq_similarity = SequenceMatcher(None, target_sentence, sentence).ratio()
            
            # 2. 关键词重叠度
            sent_keywords = set(jieba.cut(sentence))
            if len(target_keywords) > 0:
                keyword_overlap = len(target_keywords & sent_keywords) / len(target_keywords)
            else:
                keyword_overlap = 0
            
            # 3. 长度相似度
            len_similarity = 1 - abs(len(target_sentence) - len(sentence)) / max(len(target_sentence), len(sentence))
            
            # 综合评分（加权平均）
            score = seq_similarity * 0.4 + keyword_overlap * 0.4 + len_similarity * 0.2
            
            if score > best_score:
                best_score = score
                best_match = sentence
        
        # 如果相似度较高，返回匹配结果；否则返回原句
        if best_score > 0.5:  # 阈值可调整
            return best_match, match_method
        else:
            # 如果没有好的匹配，尝试部分匹配
            for sentence in sentences:
                if any(keyword in sentence for keyword in target_keywords if len(keyword) > 1):
                    return sentence, 'partial'
            
            # 实在找不到，返回原句
            return target_sentence, 'none'
    
    def _split_sentences(self, text):
        """分句"""
        # 按标点符号分句
        sentences = re.split(r'[。！？\n]+|[.!?]+\s', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 5]
        return sentences
    
    def _extract_title(self, sentences):
        """提取标题（第一句或最短的句子）"""
        if not sentences:
            return "未命名文档"
        
        # 如果第一句较短，作为标题
        if len(sentences[0]) < 50:
            return sentences[0]
        
        # 否则找最短的句子
        short_sentences = [s for s in sentences[:5] if len(s) < 50]
        if short_sentences:
            return short_sentences[0]
        
        return sentences[0][:50] + "..."
    
    def _get_color_by_importance(self, score):
        """根据重要性返回颜色"""
        if score >= 0.7:
            return '#ff6b6b'  # 红色 - 极其重要
        elif score >= 0.5:
            return '#ffa94d'  # 橙色 - 很重要
        elif score >= 0.3:
            return '#ffd43b'  # 黄色 - 重要
        else:
            return '#74c0fc'  # 蓝色 - 一般重要
    
    def _generate_statistics(self, text, keypoints):
        """生成文本统计信息"""
        word_count = len(text)
        keypoint_count = len(keypoints)
        keywords = jieba.analyse.extract_tags(text, topK=10)
        
        return {
            'word_count': word_count,
            'keypoint_count': keypoint_count,
            'top_keywords': keywords
        }
    
    def _generate_highlights(self, text, keypoints):
        """生成高亮标注元数据"""
        highlights = []
        for kp in keypoints:
            content = kp['content']
            start_pos = text.find(content)
            if start_pos != -1:
                highlights.append({
                    'id': kp['id'],
                    'start': start_pos,
                    'end': start_pos + len(content),
                    'text': content,
                    'color': kp['color'],
                    'importance': kp['importance'],
                    'annotation': kp['annotation'],
                    'category': kp['category']
                })
        return highlights

    def annotate_docx(self, input_path, highlights, output_path):
        """
        利用 Word 的 XML 结构进行物理标注
        """
        try:
            doc = Document(input_path)
            
            # 颜色映射 - 统一核心论点红色，论据蓝绿色
            color_map = {
                '#ff6b6b': WD_COLOR_INDEX.RED,         # 核心论点 - 统一红色
                '#51cf66': WD_COLOR_INDEX.BRIGHT_GREEN,# 支撑论据 - 统一亮绿色
                '#ffa94d': WD_COLOR_INDEX.DARK_YELLOW, # 兼容旧数据
                '#ffd43b': WD_COLOR_INDEX.YELLOW,      # 兼容旧数据
                '#74c0fc': WD_COLOR_INDEX.TURQUOISE    # 兼容旧数据
            }

            for hl in highlights:
                text_to_find = hl['text'].strip()
                if not text_to_find: continue
                
                # 遍历段落搜索
                for paragraph in doc.paragraphs:
                    if text_to_find in paragraph.text:
                        # 简单的替换式高亮（注意：这会破坏原文复杂的格式，但在分析预览中很有效）
                        # 对于精细标注，我们需要在 XML 层级分割 Run，这里采用稳健的全局标记法
                        # 将匹配到的文本拆分并重新设置 Run
                        if text_to_find in paragraph.text:
                            # 确定颜色
                            docx_color = color_map.get(hl['color'], WD_COLOR_INDEX.YELLOW)
                            
                            # 逻辑：将段落内容按匹配项拆分，重新构建 Runs
                            # 这种方法能确保 XML 结构中正确插入 <w:highlight>
                            orig_text = paragraph.text
                            parts = orig_text.split(text_to_find)
                            
                            paragraph.clear() # 清空旧 Runs
                            for i, part in enumerate(parts):
                                if part: paragraph.add_run(part)
                                if i < len(parts) - 1:
                                    run = paragraph.add_run(text_to_find)
                                    run.font.highlight_color = docx_color
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Docx 标注失败: {str(e)}")
            return False

    def annotate_pdf(self, input_path, highlights, output_path):
        """
        PyMuPDF 单引擎 PDF 标注（所见即所搜）

        使用 page.search_for() 搜索文本 + page.add_highlight_annot() 添加原生高亮，
        与解析阶段的 page.get_text() 共享同一文本源，彻底消除文本源断裂。

        返回: [{'id': hl_id, 'page': page_num}, ...]
        """
        try:
            doc = fitz.open(input_path)
            n_pages = len(doc)
            meta_map = []
            located = 0

            # 检测 PDF 类型（辅助调试）
            pdf_type = self._detect_pdf_type(doc)
            print(f"\U0001f4c4 PDF 类型检测: {pdf_type} (producer={doc.metadata.get('producer','')}, creator={doc.metadata.get('creator','')})")

            for hl in highlights:
                text = hl['text'].strip()
                hl_id = hl['id']
                page_hint = hl.get('page')

                if not text:
                    meta_map.append({'id': hl_id, 'page': page_hint or 1})
                    continue

                # 解析颜色
                hex_col = hl['color'].lstrip('#')
                color = (
                    int(hex_col[0:2], 16) / 255.0,
                    int(hex_col[2:4], 16) / 255.0,
                    int(hex_col[4:6], 16) / 255.0
                )

                # 确定搜索页面顺序（优先提示页附近 → 其余页面）
                if page_hint and 1 <= page_hint <= n_pages:
                    search_pages = [page_hint - 1]
                    if page_hint > 1:
                        search_pages.append(page_hint - 2)
                    if page_hint < n_pages:
                        search_pages.append(page_hint)
                    remaining = [i for i in range(n_pages) if i not in search_pages]
                else:
                    search_pages = list(range(n_pages))
                    remaining = []

                found = False
                for page_idx in search_pages + remaining:
                    found = self._try_highlight_on_page(doc[page_idx], text, color)
                    if found:
                        meta_map.append({'id': hl_id, 'page': page_idx + 1})
                        located += 1
                        print(f"  ✅ 第{page_idx + 1}页: {text[:30]}...")
                        break

                if not found:
                    meta_map.append({'id': hl_id, 'page': page_hint or 1})
                    print(f"  ⚠️ 未定位: {text[:30]}...")

            doc.save(output_path, garbage=4, deflate=True)
            doc.close()

            print(f"✅ 标注完成：{located}/{len(highlights)} 处成功定位")
            return meta_map

        except Exception as e:
            print(f"PDF 标注失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return []

    def _normalize_for_search(self, text):
        """统一文本归一化：NFKC + 去零宽字符 + 空白归一化 + 去换行"""
        t = unicodedata.normalize('NFKC', text)
        t = re.sub(r'[\u200b\u200c\u200d\ufeff\u00ad\u200e\u200f]', '', t)
        t = t.replace('\n', ' ').replace('\r', '')
        t = re.sub(r'\s+', ' ', t).strip()
        return t

    def _detect_pdf_type(self, doc):
        """
        检测 PDF 生成工具类型，用于针对性文本处理策略
        返回: 'latex' | 'word' | 'wps' | 'markdown' | 'chrome' | 'generic'
        """
        producer = (doc.metadata.get('producer', '') or '').lower()
        creator = (doc.metadata.get('creator', '') or '').lower()
        combined = producer + ' ' + creator

        if any(k in combined for k in ['latex', 'tex', 'pdftex', 'xetex', 'luatex']):
            return 'latex'
        if any(k in combined for k in ['microsoft word', 'ms word', 'docx']):
            return 'word'
        if 'wps' in combined:
            return 'wps'
        if any(k in combined for k in ['markdown', 'pandoc', 'typora']):
            return 'markdown'
        if any(k in combined for k in ['chrome', 'chromium', 'headless']):
            return 'chrome'
        return 'generic'

    def _try_highlight_on_page(self, page, text, color):
        """
        多策略搜索 + 原生高亮标注（10 级递进策略）

        针对不同 PDF 底层结构的文本差异进行适配：
        LaTeX PDF: 连字符(ﬁ→fi)、特殊排版符号
        Word PDF: 智能引号、特殊破折号
        Markdown PDF: 多余空白、代码块格式差异
        通用: 零宽字符、跨行文本、空白差异
        """
        # ── 策略 1：原始文本直接搜索 ──
        quads = page.search_for(text, quads=True)
        if quads:
            return self._apply_highlight(page, quads, color)

        # 获取页面原始文本用于后续匹配
        page_text = page.get_text("text")

        # ── 策略 2：全量归一化搜索（NFKC + 去零宽 + 去换行 + 空白归一化）──
        norm_text = self._normalize_for_search(text)
        if norm_text != text:
            quads = page.search_for(norm_text, quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)

        # ── 策略 3：智能引号/破折号替换（Word PDF 常见）──
        smart_text = norm_text
        smart_text = smart_text.replace('\u2018', "'").replace('\u2019', "'")
        smart_text = smart_text.replace('\u201c', '"').replace('\u201d', '"')
        smart_text = smart_text.replace('\u2013', '-').replace('\u2014', '-')
        smart_text = smart_text.replace('\u00a0', ' ')  # non-breaking space
        if smart_text != norm_text:
            quads = page.search_for(smart_text, quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)

        # ── 策略 4：PyMuPDF 页面文本子串匹配 → 提取精确片段搜索 ──
        norm_page = self._normalize_for_search(page_text)
        norm_query = self._normalize_for_search(text)
        if norm_query in norm_page:
            # 在归一化页面文本中找到了，用逐步截短的方式在原始PDF中搜索
            for seg_len in [len(norm_query), 80, 50, 30, 20]:
                if seg_len > len(norm_query):
                    continue
                snippet = norm_query[:seg_len]
                quads = page.search_for(snippet, quads=True)
                if quads:
                    return self._apply_highlight(page, quads, color)

        # ── 策略 5：中文标点变体（全角↔半角）──
        punct_map = str.maketrans('，。！？；：（）', ',.!?;:()')
        punct_text = norm_text.translate(punct_map)
        if punct_text != norm_text:
            quads = page.search_for(punct_text, quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)

        # ── 策略 6：Docling 与 PyMuPDF 文本差异对齐 ──
        # Docling 提取的文本可能有额外空格（如中英文间），而 PyMuPDF 没有
        # 尝试去除所有空格后搜索
        no_space = re.sub(r'\s', '', norm_query)
        page_no_space = re.sub(r'\s', '', page_text)
        if len(no_space) >= 10 and no_space in page_no_space:
            # 找到无空格匹配，用前若干字搜索
            for seg_len in [25, 18, 12, 8]:
                if seg_len > len(norm_query):
                    continue
                snippet = norm_query[:seg_len]
                quads = page.search_for(snippet, quads=True)
                if quads:
                    return self._apply_highlight(page, quads, color)

        # ── 策略 7：逐级截短搜索（前段 + 后段 + 中段）──
        for length in [80, 60, 40, 25, 15]:
            if len(norm_text) <= length:
                continue
            # 前段
            quads = page.search_for(norm_text[:length], quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)
            # 后段
            quads = page.search_for(norm_text[-length:], quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)
            # 中段
            mid = len(norm_text) // 2
            mid_start = max(0, mid - length // 2)
            quads = page.search_for(norm_text[mid_start:mid_start + length], quads=True)
            if quads:
                return self._apply_highlight(page, quads, color)

        # ── 策略 8：关键词锚定搜索（提取长词在页面中定位）──
        keywords = [w for w in jieba.cut(text) if len(w) >= 4]
        if keywords:
            # 用最长的关键词尝试定位
            keywords.sort(key=len, reverse=True)
            for kw in keywords[:3]:
                quads = page.search_for(kw, quads=True)
                if quads:
                    return self._apply_highlight(page, quads, color)

        # ── 策略 9：滑动窗口搜索（在页面文本中寻找最佳匹配子串）──
        if len(norm_query) >= 15:
            best_ratio = 0
            best_substr = None
            window_size = min(len(norm_query) + 20, len(norm_page))
            step = max(1, window_size // 4)
            for i in range(0, max(1, len(norm_page) - window_size + 1), step):
                window = norm_page[i:i + window_size]
                ratio = SequenceMatcher(None, norm_query, window).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_substr = window
            if best_ratio > 0.7 and best_substr:
                # 找到高相似度窗口，用前20字在原始PDF中搜索
                for seg_len in [25, 18, 12]:
                    snippet = best_substr[:seg_len]
                    quads = page.search_for(snippet, quads=True)
                    if quads:
                        return self._apply_highlight(page, quads, color)

        return False

    def _apply_highlight(self, page, quads, color):
        """应用高亮标注"""
        annot = page.add_highlight_annot(quads)
        annot.set_colors(stroke=color)
        annot.set_opacity(0.35)
        annot.update()
        return True
    
    
    def compare_documents(self, records):
        """比较多个文档"""
        comparison = {
            'common_keywords': [],
            'unique_keywords': {},
            'similarity': {}
        }
        
        # 提取每个文档的关键词
        all_keywords = {}
        for record in records:
            if record and 'analysis_result' in record:
                keywords = record['analysis_result']['statistics']['top_keywords']
                all_keywords[record['id']] = set(keywords)
        
        # 找出共同关键词
        if len(all_keywords) >= 2:
            common = set.intersection(*all_keywords.values())
            comparison['common_keywords'] = list(common)
            
            # 找出每个文档的独特关键词
            for record_id, keywords in all_keywords.items():
                unique = keywords - set(comparison['common_keywords'])
                comparison['unique_keywords'][record_id] = list(unique)
        
        return comparison
    
    def _locate_text_in_pdf(self, pdf_path, text, page_hint=None):
        """
        使用 PyMuPDF search_for 在 PDF 中精准定位文本位置
        增强版：多种文本归一化变体 + 逐级截短降级
        返回: (page_num, bbox_dict) 或 None
        """
        if not pdf_path or not os.path.exists(pdf_path):
            return None

        try:
            doc = fitz.open(pdf_path)
            n_pages = len(doc)

            # 决定搜索范围（优先提示页附近）
            if page_hint and 1 <= page_hint <= n_pages:
                search_pages = [page_hint - 1]
                for offset in [1, -1, 2, -2]:
                    p = page_hint - 1 + offset
                    if 0 <= p < n_pages and p not in search_pages:
                        search_pages.append(p)
                remaining = [i for i in range(n_pages) if i not in search_pages]
                search_pages += remaining
            else:
                search_pages = list(range(n_pages))

            # 准备多种搜索变体
            text_clean = text.strip()
            norm = self._normalize_for_search(text_clean)
            
            # 智能引号/破折号替换
            smart = norm.replace('\u2018', "'").replace('\u2019', "'")
            smart = smart.replace('\u201c', '"').replace('\u201d', '"')
            smart = smart.replace('\u2013', '-').replace('\u2014', '-')
            smart = smart.replace('\u00a0', ' ')

            variants = [text_clean]
            for v in [norm, smart]:
                if v not in variants:
                    variants.append(v)

            for page_idx in search_pages:
                page = doc[page_idx]

                # 尝试所有变体
                for variant in variants:
                    rects = page.search_for(variant)
                    if rects:
                        doc.close()
                        return (page_idx + 1, {
                            'x0': rects[0].x0, 'y0': rects[0].y0,
                            'x1': rects[0].x1, 'y1': rects[0].y1
                        })

                # 归一化子串匹配
                page_text = page.get_text("text")
                norm_page = self._normalize_for_search(page_text)
                if norm in norm_page:
                    for seg_len in [30, 20, 12]:
                        if seg_len > len(norm):
                            continue
                        rects = page.search_for(norm[:seg_len])
                        if rects:
                            doc.close()
                            return (page_idx + 1, {
                                'x0': rects[0].x0, 'y0': rects[0].y0,
                                'x1': rects[0].x1, 'y1': rects[0].y1
                            })

                # 截短降级
                for length in [40, 25, 15]:
                    if len(text_clean) <= length:
                        continue
                    for variant in variants:
                        if len(variant) > length:
                            rects = page.search_for(variant[:length])
                            if rects:
                                doc.close()
                                return (page_idx + 1, {
                                    'x0': rects[0].x0, 'y0': rects[0].y0,
                                    'x1': rects[0].x1, 'y1': rects[0].y1
                                })

            doc.close()
            return None

        except Exception as e:
            print(f"PDF定位错误: {str(e)}")
            return None
