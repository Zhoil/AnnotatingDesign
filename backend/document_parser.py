import os
from bs4 import BeautifulSoup
from docx import Document
import re
import markdown
from pdf_parser_enhanced import EnhancedPDFParser
from web_parser_enhanced import EnhancedWebParser

class DocumentParser:
    """
    文档解析器
    支持格式: PDF, Word, HTML, Markdown, 网页URL
    PDF 使用增强版解析器，支持表格和图片提取
    """
    
    def __init__(self):
        self.enhanced_pdf_parser = EnhancedPDFParser()
        self.enhanced_web_parser = EnhancedWebParser()
    
    def parse(self, filepath):
        """
        根据文件类型解析文档
        返回：{'text': str, 'structured_content': list, 'metadata': dict, 'file_size': int, 'filepath': str}
        """
        ext = os.path.splitext(filepath)[1].lower()
        file_size = os.path.getsize(filepath)
            
        if ext == '.pdf':
            # PDF: 使用增强版解析器（支持表格、图片）
            result = self.enhanced_pdf_parser.parse_pdf_enhanced(filepath)
            if result:
                result['file_size'] = file_size
                result['filepath'] = filepath
                return result
        elif ext == '.docx':
            return {
                'text': self.parse_docx(filepath),
                'structured_content': [],
                'metadata': {},
                'file_size': file_size,
                'filepath': filepath
            }
        elif ext == '.doc':
            # .doc 是旧版 OLE 二进制格式，python-docx 不支持，需用 win32com 调用 Word
            return {
                'text': self.parse_doc_legacy(filepath),
                'structured_content': [],
                'metadata': {},
                'file_size': file_size,
                'filepath': filepath
            }
        elif ext in ['.html', '.htm']:
            return {
                'text': self.parse_html_file(filepath), 
                'structured_content': [], 
                'metadata': {},
                'file_size': file_size,
                'filepath': filepath
            }
        elif ext in ['.md', '.markdown']:
            return {
                'text': self.parse_markdown(filepath), 
                'structured_content': [], 
                'metadata': {},
                'file_size': file_size,
                'filepath': filepath
            }
        else:
            raise ValueError(f"不支持的文件格式：{ext}")
        
    def parse_docx(self, filepath):
        """解析 .docx 文件（python-docx）"""
        try:
            doc = Document(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return self._clean_text(text)
        except Exception as e:
            print(f"Word(.docx)解析错误: {str(e)}")
            return None

    def parse_doc_legacy(self, filepath):
        """
        解析旧版 .doc 文件（OLE 二进制格式）。依次尝试：
          1. 子进程隔离的 win32com（解决 Flask 多线程 COM 冲突）
          2. 当前线程直接 win32com（降级备用）
          3. 二进制文本提取（完全不需要 Word）
        """
        print(f"📂 尝试解析 .doc 文件: {os.path.basename(filepath)}")

        # ─ 方法 1：子进程隔离（最可靠）─────────────────────
        text = self._parse_doc_subprocess(filepath)
        if text:
            print("✅ .doc 子进程解析成功")
            return text

        # ─ 方法 2：当前线程直接 win32com ────────────────────
        text = self._parse_doc_win32com(filepath)
        if text:
            print("✅ .doc win32com 解析成功")
            return text

        # ─ 方法 3：二进制提取（最后备用）───────────────────
        text = self._parse_doc_binary(filepath)
        if text:
            print("⚠️ .doc 二进制提取成功（质量有限）")
            return text

        print("❌ .doc 所有解析方法均失败")
        return None

    def _parse_doc_subprocess(self, filepath):
        """
        子进程隔离执行 win32com。
        Flask 多线程环境下，COM（STA）和工作线程存在冲突。
        独立子进程不存在这个问题。
        """
        import subprocess
        import sys
        import tempfile
        import uuid

        abs_path = os.path.abspath(filepath)
        tmp_out  = os.path.join(tempfile.gettempdir(),
                                f'doc_text_{uuid.uuid4().hex}.txt')

        # 内联脚本：完整运行在子进程的主线程中
        script = (
            'import sys, pythoncom, win32com.client\n'
            'pythoncom.CoInitialize()\n'
            'word = None\n'
            'try:\n'
            '    word = win32com.client.Dispatch("Word.Application")\n'
            '    word.Visible = False\n'
            '    word.DisplayAlerts = 0\n'
            '    word.AutomationSecurity = 3\n'
            f'    doc = word.Documents.Open(r"{abs_path}", ReadOnly=True,'
            ' AddToRecentFiles=False, NoEncodingDialog=True,'
            ' ConfirmConversions=False)\n'
            f'    open(r"{tmp_out}", "w", encoding="utf-8").write(doc.Content.Text)\n'
            '    doc.Close(False)\n'
            'except Exception as e:\n'
            '    print(f"ERR:{e}", file=sys.stderr)\n'
            '    sys.exit(1)\n'
            'finally:\n'
            '    if word:\n'
            '        try: word.Quit()\n'
            '        except: pass\n'
            '    pythoncom.CoUninitialize()\n'
        )

        try:
            result = subprocess.run(
                [sys.executable, '-c', script],
                capture_output=True, text=True, timeout=45
            )
            if result.returncode == 0 and os.path.exists(tmp_out):
                with open(tmp_out, 'r', encoding='utf-8') as f:
                    text = f.read()
                if text.strip():
                    return self._clean_text(text)
            if result.stderr:
                print(f"  Word 子进程: {result.stderr[:200].strip()}")
        except subprocess.TimeoutExpired:
            print("  Word 子进程超时")
        except Exception as e:
            print(f"  子进程执行失败: {e}")
        finally:
            try:
                if os.path.exists(tmp_out):
                    os.unlink(tmp_out)
            except Exception:
                pass
        return None

    def _parse_doc_win32com(self, filepath):
        """win32com 直接调用（当前线程，备用）"""
        word = None
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            word.AutomationSecurity = 3
            abs_path = os.path.abspath(filepath)
            doc = word.Documents.Open(
                abs_path,
                ReadOnly=True, AddToRecentFiles=False,
                NoEncodingDialog=True, ConfirmConversions=False
            )
            text = doc.Content.Text
            doc.Close(False)
            return self._clean_text(text) if text.strip() else None
        except ImportError:
            print("  win32com 不可用")
            return None
        except Exception as e:
            print(f"  win32com 直接调用失败: {e}")
            return None
        finally:
            try:
                if word:
                    word.Quit()
            except Exception:
                pass
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _parse_doc_binary(self, filepath):
        """
        二进制提取法——完全不需要 Word。
        支持两种情况：
          a) 文件实际是 OOXML（重命名为 .doc）→ 用 python-docx 读
          b) 真正的 OLE 二进制 → 提取 UTF-16 LE 文本块
        """
        try:
            with open(filepath, 'rb') as f:
                data = f.read()

            OLE_MAGIC = b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'

            # 情况 a：实际是 ZIP（OOXML 文件被重命名为 .doc）
            if data[:4] == b'PK\x03\x04':
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document(BytesIO(data))
                    text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                    if text.strip():
                        print("  检测到该 .doc 实际是 OOXML 格式")
                        return self._clean_text(text)
                except Exception as e:
                    print(f"  OOXML 尝试失败: {e}")

            # 情况 b：标准 OLE 二进制
            if data[:8] == OLE_MAGIC:
                # Word 97-2003 以 UTF-16 LE 存储文本
                try:
                    raw = data.decode('utf-16-le', errors='ignore')
                    # 过滤控制字符，保留中英文数及常用标点
                    clean = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', raw)
                    lines = [l.strip() for l in re.split(r'[\r\n\x0b\x0c]+', clean)
                             if len(l.strip()) > 4]
                    text = '\n'.join(lines)
                    if len(text) > 80:
                        return self._clean_text(text)
                except Exception as e:
                    print(f"  UTF-16 提取失败: {e}")

        except Exception as e:
            print(f"  二进制提取异常: {e}")

        return None

    
    def parse_html_file(self, filepath):
        """解析HTML文件"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            return self._clean_text(text)
        except Exception as e:
            print(f"HTML解析错误: {str(e)}")
            return None
    
    def parse_url(self, url):
        """解析网页 URL，使用增强版解析器保留 HTML 结构"""
        return self.enhanced_web_parser.parse_url(url)
    
    def parse_markdown(self, filepath):
        """解析Markdown文件"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # 将Markdown转换为HTML
            html = markdown.markdown(md_content)
            
            # 使用BeautifulSoup提取纯文本
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            return self._clean_text(text)
        except Exception as e:
            print(f"Markdown解析错误: {str(e)}")
            return None
    
    def _clean_text(self, text):
        """清理文本"""
        # 移除多余的空白字符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = text.strip()
        return text
